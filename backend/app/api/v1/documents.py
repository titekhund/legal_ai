"""
Document Generation API endpoints
"""
import io
import time
import uuid
from collections import OrderedDict
from datetime import datetime, timedelta
from typing import Dict, List, Optional

from fastapi import APIRouter, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.core import get_logger
from app.models import (
    DocumentGenerationRequest,
    DocumentTemplate,
    DocumentType,
    GeneratedDocument,
)
from app.services import DocumentService

logger = get_logger(__name__)
router = APIRouter()

# Global service instance (will be set from main.py)
_document_service: Optional[DocumentService] = None


# ============================================================================
# Document Storage (In-Memory with TTL)
# ============================================================================


class StoredDocument:
    """Document stored in memory with metadata"""

    def __init__(
        self,
        document: GeneratedDocument,
        document_id: str,
        expires_at: datetime
    ):
        self.document = document
        self.document_id = document_id
        self.expires_at = expires_at
        self.created_at = datetime.utcnow()


class DocumentStore:
    """
    In-memory document storage with TTL and size limits
    """

    def __init__(self, max_size: int = 100, ttl_seconds: int = 3600):
        """
        Initialize document store

        Args:
            max_size: Maximum number of documents to store
            ttl_seconds: Time to live for documents (default 1 hour)
        """
        self.documents: OrderedDict[str, StoredDocument] = OrderedDict()
        self.max_size = max_size
        self.ttl_seconds = ttl_seconds

    def store(self, document: GeneratedDocument) -> str:
        """
        Store a document and return its ID

        Args:
            document: Generated document to store

        Returns:
            Document ID for retrieval
        """
        # Clean up expired documents
        self._cleanup_expired()

        # Generate unique ID
        document_id = str(uuid.uuid4())

        # Calculate expiration
        expires_at = datetime.utcnow() + timedelta(seconds=self.ttl_seconds)

        # Store document
        stored_doc = StoredDocument(
            document=document,
            document_id=document_id,
            expires_at=expires_at
        )
        self.documents[document_id] = stored_doc

        # Enforce size limit (FIFO)
        if len(self.documents) > self.max_size:
            # Remove oldest document
            self.documents.popitem(last=False)
            logger.info("Removed oldest document due to size limit")

        logger.info(f"Stored document: {document_id}, expires: {expires_at}")
        return document_id

    def retrieve(self, document_id: str) -> Optional[GeneratedDocument]:
        """
        Retrieve a stored document

        Args:
            document_id: Document ID

        Returns:
            GeneratedDocument or None if not found/expired
        """
        stored_doc = self.documents.get(document_id)

        if not stored_doc:
            return None

        # Check if expired
        if datetime.utcnow() > stored_doc.expires_at:
            del self.documents[document_id]
            logger.info(f"Document expired: {document_id}")
            return None

        return stored_doc.document

    def _cleanup_expired(self):
        """Remove expired documents"""
        now = datetime.utcnow()
        expired_ids = [
            doc_id
            for doc_id, stored_doc in self.documents.items()
            if now > stored_doc.expires_at
        ]

        for doc_id in expired_ids:
            del self.documents[doc_id]

        if expired_ids:
            logger.info(f"Cleaned up {len(expired_ids)} expired documents")

    def get_stats(self) -> dict:
        """Get storage statistics"""
        self._cleanup_expired()
        return {
            "total_documents": len(self.documents),
            "max_size": self.max_size,
            "ttl_seconds": self.ttl_seconds
        }


# Global document store
_document_store = DocumentStore(max_size=100, ttl_seconds=3600)


def set_document_service(service: DocumentService):
    """Set the document service instance"""
    global _document_service
    _document_service = service


def get_document_service() -> DocumentService:
    """Get the document service instance"""
    if _document_service is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Document service not initialized"
        )
    return _document_service


# ============================================================================
# Request/Response Models
# ============================================================================


class TemplateSearchResponse(BaseModel):
    """Template search response"""

    templates: List[DocumentTemplate]
    total: int


class DocumentTypesResponse(BaseModel):
    """Document types response"""

    types: List[DocumentType]
    total: int


class DownloadLinks(BaseModel):
    """Download links for generated document"""

    markdown: str
    docx: str
    pdf: Optional[str] = None


class DocumentGenerationResponse(BaseModel):
    """Document generation response"""

    content: str
    document_type: str
    template_used: str
    cited_articles: List[str]
    variables_used: Dict
    disclaimer: str
    format: str
    warnings: List[str]
    processing_time_ms: int
    download_links: DownloadLinks


# ============================================================================
# Format Conversion
# ============================================================================


def markdown_to_docx(markdown_content: str) -> bytes:
    """
    Convert markdown to DOCX format

    Args:
        markdown_content: Markdown formatted text

    Returns:
        DOCX file bytes
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re

        doc = Document()

        # Process markdown line by line
        lines = markdown_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # Headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
            # Bold text
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            # Horizontal rule
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            # Lists
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
            # Regular paragraph
            else:
                # Handle inline formatting
                p = doc.add_paragraph()
                text = line

                # Simple bold conversion
                parts = re.split(r'\*\*(.*?)\*\*', text)
                for j, part in enumerate(parts):
                    if j % 2 == 0:
                        p.add_run(part)
                    else:
                        run = p.add_run(part)
                        run.bold = True

            i += 1

        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        return docx_bytes.getvalue()

    except ImportError:
        logger.error("python-docx not installed")
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="DOCX conversion not available - python-docx not installed"
        )
    except Exception as e:
        logger.error(f"Error converting to DOCX: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to convert to DOCX: {str(e)}"
        )


# ============================================================================
# Endpoints
# ============================================================================


@router.get("/documents/types", response_model=DocumentTypesResponse)
async def list_document_types(language: str = "ka"):
    """
    List all available document types

    Args:
        language: Language filter ('ka' or 'en')

    Returns:
        List of document types with descriptions
    """
    try:
        service = get_document_service()
        types = await service.list_document_types(language=language)

        logger.info(f"Listed {len(types)} document types", extra={"language": language})

        return DocumentTypesResponse(types=types, total=len(types))

    except Exception as e:
        logger.error(f"Error listing document types: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list document types: {str(e)}"
        )


@router.get("/documents/types/{type_id}", response_model=DocumentType)
async def get_document_type(type_id: str):
    """
    Get a specific document type by ID

    Args:
        type_id: Document type identifier

    Returns:
        Document type with full field definitions
    """
    try:
        service = get_document_service()
        template_store = service.template_store

        doc_type = template_store.get_document_type(type_id)

        if not doc_type:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document type not found: {type_id}"
            )

        logger.info(f"Retrieved document type: {type_id}")
        return doc_type

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document type: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get document type: {str(e)}"
        )


@router.get("/documents/templates", response_model=TemplateSearchResponse)
async def search_templates(
    query: Optional[str] = None,
    document_type: Optional[str] = None,
    language: str = "ka",
    limit: int = 10
):
    """
    Search for document templates

    Args:
        query: Search query (searches in name and category)
        document_type: Filter by document type (type parameter)
        language: Language filter ('ka' or 'en')
        limit: Maximum number of results

    Returns:
        List of matching templates
    """
    try:
        service = get_document_service()

        if query:
            templates = await service.search_templates(
                query=query, document_type=document_type, language=language
            )
        else:
            # If no query, get all templates of the specified type
            templates = await service.search_templates(
                query="", language=language
            )

            # Filter by document type if specified
            if document_type:
                templates = [t for t in templates if t.type == document_type]

        # Apply limit
        templates = templates[:limit]

        logger.info(
            f"Found {len(templates)} templates",
            extra={"query": query, "document_type": document_type, "language": language}
        )

        return TemplateSearchResponse(templates=templates, total=len(templates))

    except Exception as e:
        logger.error(f"Error searching templates: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to search templates: {str(e)}"
        )


@router.get("/documents/templates/{template_id}", response_model=DocumentTemplate)
async def get_template(template_id: str):
    """
    Get a specific template by ID

    Args:
        template_id: Template identifier

    Returns:
        Template details including variables and content
    """
    try:
        service = get_document_service()
        template_store = service.template_store

        if not template_store or template_id not in template_store.templates:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Template not found: {template_id}"
            )

        template = template_store.templates[template_id]
        logger.info(f"Retrieved template: {template_id}")
        return template

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting template: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get template: {str(e)}"
        )


@router.post("/documents/generate", response_model=DocumentGenerationResponse)
async def generate_document(request: DocumentGenerationRequest):
    """
    Generate a document from a template

    Args:
        request: Document generation request with template ID and variables

    Returns:
        Generated document with content, metadata, and download links
    """
    start_time = time.time()

    try:
        service = get_document_service()

        logger.info(
            f"Generating document from template: {request.template_id}",
            extra={
                "template_id": request.template_id,
                "document_type": request.document_type,
                "format": request.format
            }
        )

        # Generate document
        document = await service.generate_document(request)

        # Store document for downloads
        document_id = _document_store.store(document)

        # Create download links
        base_url = "/v1/documents/download"
        download_links = DownloadLinks(
            markdown=f"{base_url}/{document_id}?format=md",
            docx=f"{base_url}/{document_id}?format=docx",
            pdf=None  # PDF not implemented in MVP
        )

        processing_time_ms = int((time.time() - start_time) * 1000)

        logger.info(
            f"Document generated successfully: {document_id}",
            extra={
                "document_id": document_id,
                "template_id": request.template_id,
                "format": document.format,
                "processing_time_ms": processing_time_ms,
                "warnings": len(document.warnings)
            }
        )

        return DocumentGenerationResponse(
            content=document.content,
            document_type=document.document_type,
            template_used=document.template_used,
            cited_articles=document.cited_articles,
            variables_used=document.variables_used,
            disclaimer=document.disclaimer,
            format=document.format,
            warnings=document.warnings,
            processing_time_ms=processing_time_ms,
            download_links=download_links
        )

    except ValueError as e:
        logger.warning(f"Validation error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        logger.error(f"Error generating document: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate document: {str(e)}"
        )


@router.get("/documents/download/{document_id}")
async def download_document(
    document_id: str,
    format: str = Query("md", description="Download format: md, docx, pdf")
):
    """
    Download a generated document in specified format

    Args:
        document_id: Document identifier from generation response
        format: Output format (md, docx, pdf)

    Returns:
        File download stream
    """
    try:
        # Retrieve document from storage
        document = _document_store.retrieve(document_id)

        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found or expired"
            )

        # Determine filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{document.document_type}_{timestamp}"

        # Convert and return based on format
        if format == "md" or format == "markdown":
            content = document.content.encode('utf-8')
            media_type = "text/markdown"
            filename = f"{base_filename}.md"

        elif format == "docx":
            content = markdown_to_docx(document.content)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{base_filename}.docx"

        elif format == "pdf":
            # PDF conversion not implemented in MVP
            raise HTTPException(
                status_code=status.HTTP_501_NOT_IMPLEMENTED,
                detail="PDF format not yet implemented"
            )

        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported format: {format}. Use 'md' or 'docx'"
            )

        logger.info(
            f"Downloaded document: {document_id}",
            extra={"format": format, "filename": filename}
        )

        return StreamingResponse(
            io.BytesIO(content),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading document: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to download document: {str(e)}"
        )


@router.get("/documents/storage/stats")
async def get_storage_stats():
    """
    Get document storage statistics (admin/debug endpoint)

    Returns:
        Storage statistics
    """
    return _document_store.get_stats()
